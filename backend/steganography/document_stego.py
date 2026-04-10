"""
Document Steganography Handler
==============================
Supports: TXT, PDF, DOCX formats
Methods: Zero-width characters, metadata injection, stream modification
"""

import os
import io
import json
import zipfile
from typing import Optional, Tuple, Dict, Any
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TXTSteganography:
    """Text file steganography using various encoding methods"""
    
    # Unicode zero-width characters
    ZERO_WIDTH_SPACE = '\u200B'      # Zero-Width Space
    ZERO_WIDTH_NON_JOINER = '\u200C' # Zero-Width Non-Joiner
    ZERO_WIDTH_JOINER = '\u200D'     # Zero-Width Joiner
    
    @staticmethod
    def hide_zero_width(secret_data: bytes, text_content: str, output_file: str) -> Dict[str, Any]:
        """
        Hide secret data using zero-width Unicode characters
        
        Encoding:
        - 0 bit: U+200B (Zero-Width Space)
        - 1 bit: U+200C (Zero-Width Non-Joiner)
        - Separator: U+200D (Zero-Width Joiner)
        """
        if len(secret_data) == 0:
            raise ValueError("Secret data cannot be empty")
        
        # Convert bytes to binary string
        binary_str = ''.join(f'{byte:08b}' for byte in secret_data)
        
        # Encode binary string as zero-width characters
        encoded = ''
        for bit in binary_str:
            if bit == '0':
                encoded += TXTSteganography.ZERO_WIDTH_SPACE
            else:
                encoded += TXTSteganography.ZERO_WIDTH_NON_JOINER
        
        # Add separator at end
        encoded += TXTSteganography.ZERO_WIDTH_JOINER
        
        # Insert into text (after first character to ensure it's not stripped)
        if len(text_content) == 0:
            text_content = 'Document'
        
        stego_content = text_content[0] + encoded + text_content[1:]
        
        # Write to file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(stego_content)
        
        return {
            'method': 'zero-width',
            'secret_size': len(secret_data),
            'text_size': len(text_content),
            'capacity_used': f"{(len(secret_data) / len(text_content)) * 100:.2f}%"
        }
    
    @staticmethod
    def extract_zero_width(stego_file: str) -> bytes:
        """Extract data hidden with zero-width characters"""
        with open(stego_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Find zero-width character sequence
        binary_str = ''
        i = 1  # Skip first character
        while i < len(content):
            char = content[i]
            if char == TXTSteganography.ZERO_WIDTH_SPACE:
                binary_str += '0'
            elif char == TXTSteganography.ZERO_WIDTH_NON_JOINER:
                binary_str += '1'
            elif char == TXTSteganography.ZERO_WIDTH_JOINER:
                break  # End of hidden data
            i += 1
        
        # Convert binary string to bytes
        secret_data = bytearray()
        for i in range(0, len(binary_str), 8):
            byte_binary = binary_str[i:i+8]
            if len(byte_binary) == 8:
                secret_data.append(int(byte_binary, 2))
        
        return bytes(secret_data)


class PDFSteganography:
    """PDF steganography using metadata and stream manipulation"""
    
    @staticmethod
    def hide_in_metadata(secret_data: bytes, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """
        Hide secret data in PDF metadata dictionary
        Stores encoded data in /Producer, /Creator, /Subject fields
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            # Read PDF
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                writer = PyPDF2.PdfWriter()
                
                # Copy all pages
                for page_num in range(len(reader.pages)):
                    writer.add_page(reader.pages[page_num])
                
                # Encode secret data
                encoded_secret = PDFSteganography._encode_secret(secret_data)
                
                # Add to metadata
                metadata = reader.metadata.copy() if reader.metadata else {}
                metadata['/Producer'] = f"Stego_{encoded_secret[:50]}"
                metadata['/Creator'] = f"StegoDoc_{encoded_secret[50:100]}"
                metadata['/Subject'] = f"Hidden_{encoded_secret[100:150]}"
                
                writer.add_metadata(metadata)
                
                # Write output
                with open(output_path, 'wb') as f:
                    writer.write(f)
            
            return {
                'method': 'metadata',
                'secret_size': len(secret_data),
                'pdf_pages': len(reader.pages),
                'storage': 'PDF metadata fields'
            }
        
        except Exception as e:
            raise RuntimeError(f"PDF metadata hiding failed: {str(e)}")
    
    @staticmethod
    def extract_from_metadata(pdf_path: str) -> bytes:
        """Extract data hidden in PDF metadata"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed")
        
        try:
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata = reader.metadata
                
                if not metadata:
                    return b''
                
                # Reconstruct secret from metadata
                encoded = ''
                if '/Producer' in metadata:
                    encoded += str(metadata['/Producer'])[6:56] if len(str(metadata['/Producer'])) > 6 else ''
                if '/Creator' in metadata:
                    encoded += str(metadata['/Creator'])[8:58] if len(str(metadata['/Creator'])) > 8 else ''
                if '/Subject' in metadata:
                    encoded += str(metadata['/Subject'])[7:57] if len(str(metadata['/Subject'])) > 7 else ''
                
                return PDFSteganography._decode_secret(encoded)
        
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {str(e)}")
    
    @staticmethod
    def _encode_secret(data: bytes) -> str:
        """Encode bytes to hex string for storage"""
        return data.hex()
    
    @staticmethod
    def _decode_secret(hex_str: str) -> bytes:
        """Decode hex string back to bytes"""
        try:
            return bytes.fromhex(hex_str)
        except:
            return b''


class DOCXSteganography:
    """DOCX steganography using XML manipulation"""
    
    @staticmethod
    def hide_in_properties(secret_data: bytes, docx_path: str, output_path: str) -> Dict[str, Any]:
        """
        Hide secret data in DOCX core properties and custom XML
        Modified properties: Creator, Subject, Title, Description
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Load document
            doc = Document(docx_path)
            
            # Encode secret
            encoded_secret = DOCXSteganography._encode_secret(secret_data)
            
            # Modify core properties
            doc.core_properties.creator = f"Author_{encoded_secret[:30]}"
            doc.core_properties.subject = f"Doc_{encoded_secret[30:60]}"
            doc.core_properties.title = f"Hidden_{encoded_secret[60:90]}"
            doc.core_properties.comments = f"Data_{encoded_secret[90:120]}"
            
            # Save document
            doc.save(output_path)
            
            return {
                'method': 'core_properties',
                'secret_size': len(secret_data),
                'document_paragraphs': len(doc.paragraphs),
                'storage': 'DOCX core properties'
            }
        
        except Exception as e:
            raise RuntimeError(f"DOCX property hiding failed: {str(e)}")
    
    @staticmethod
    def extract_from_properties(docx_path: str) -> bytes:
        """Extract data hidden in DOCX properties"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        try:
            doc = Document(docx_path)
            props = doc.core_properties
            
            # Reconstruct secret from properties
            encoded = ''
            if props.creator and len(str(props.creator)) > 7:
                encoded += str(props.creator)[7:37]
            if props.subject and len(str(props.subject)) > 4:
                encoded += str(props.subject)[4:34]
            if props.title and len(str(props.title)) > 7:
                encoded += str(props.title)[7:37]
            if props.comments and len(str(props.comments)) > 5:
                encoded += str(props.comments)[5:35]
            
            return DOCXSteganography._decode_secret(encoded)
        
        except Exception as e:
            raise RuntimeError(f"DOCX extraction failed: {str(e)}")
    
    @staticmethod
    def _encode_secret(data: bytes) -> str:
        """Encode bytes to hex string for storage"""
        return data.hex()
    
    @staticmethod
    def _decode_secret(hex_str: str) -> bytes:
        """Decode hex string back to bytes"""
        try:
            # Clean hex string
            hex_str = ''.join(c for c in hex_str if c in '0123456789abcdefABCDEF')
            if len(hex_str) % 2 != 0:
                hex_str = hex_str[:-1]
            return bytes.fromhex(hex_str) if hex_str else b''
        except:
            return b''


class DocumentSteganography:
    """Unified document steganography interface"""
    
    @staticmethod
    def get_document_format(file_path: str) -> str:
        """Detect document format from file extension"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.txt':
            return 'txt'
        elif ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        else:
            raise ValueError(f"Unsupported document format: {ext}")
    
    @staticmethod
    def hide_in_document(secret_data: bytes, cover_doc: str, output_file: str, 
                        method: Optional[str] = None) -> Dict[str, Any]:
        """
        Hide data in any supported document format
        Automatically detects format and uses appropriate method
        """
        doc_format = DocumentSteganography.get_document_format(cover_doc)
        
        if doc_format == 'txt':
            with open(cover_doc, 'r', encoding='utf-8') as f:
                content = f.read()
            return TXTSteganography.hide_zero_width(secret_data, content, output_file)
        
        elif doc_format == 'pdf':
            return PDFSteganography.hide_in_metadata(secret_data, cover_doc, output_file)
        
        elif doc_format == 'docx':
            return DOCXSteganography.hide_in_properties(secret_data, cover_doc, output_file)
        
        else:
            raise ValueError(f"Unsupported document format: {doc_format}")
    
    @staticmethod
    def extract_from_document(stego_doc: str) -> bytes:
        """
        Extract data from any supported document format
        Automatically detects format and uses appropriate method
        """
        doc_format = DocumentSteganography.get_document_format(stego_doc)
        
        if doc_format == 'txt':
            return TXTSteganography.extract_zero_width(stego_doc)
        
        elif doc_format == 'pdf':
            return PDFSteganography.extract_from_metadata(stego_doc)
        
        elif doc_format == 'docx':
            return DOCXSteganography.extract_from_properties(stego_doc)
        
        else:
            raise ValueError(f"Unsupported document format: {doc_format}")
    
    @staticmethod
    def validate_capacity(secret_size: int, cover_doc: str) -> Tuple[bool, str]:
        """
        Validate if document has enough capacity for secret data
        Returns: (is_valid, message)
        """
        doc_format = DocumentSteganography.get_document_format(cover_doc)
        
        # Rough capacity estimates (in bytes)
        try:
            if doc_format == 'txt':
                with open(cover_doc, 'rb') as f:
                    cover_size = len(f.read())
                max_capacity = cover_size // 8  # ~12.5% capacity
            
            elif doc_format == 'pdf':
                with open(cover_doc, 'rb') as f:
                    cover_size = len(f.read())
                max_capacity = min(500, cover_size // 100)  # ~1% or 500 bytes max
            
            elif doc_format == 'docx':
                with open(cover_doc, 'rb') as f:
                    cover_size = len(f.read())
                max_capacity = min(1000, cover_size // 50)  # ~2% or 1KB max
            
            else:
                return False, "Unsupported format"
            
            if secret_size <= max_capacity:
                return True, f"OK - {(secret_size/max_capacity)*100:.1f}% capacity used"
            else:
                return False, f"Insufficient capacity: {secret_size} bytes needed, {max_capacity} available"
        
        except Exception as e:
            return False, f"Validation error: {str(e)}"


__all__ = [
    'TXTSteganography',
    'PDFSteganography',
    'DOCXSteganography',
    'DocumentSteganography'
]


class TXTSteganography:
    """Text file steganography using zero-width characters and whitespace"""

    # Zero-width Unicode characters
    ZERO_WIDTH_SPACE = '\u200B'      # U+200B
    ZERO_WIDTH_NON_JOINER = '\u200C'  # U+200C
    ZERO_WIDTH_JOINER = '\u200D'      # U+200D
    ZERO_WIDTH_ZERO = ZERO_WIDTH_SPACE
    ZERO_WIDTH_ONE = ZERO_WIDTH_NON_JOINER
    ZERO_WIDTH_SEP = ZERO_WIDTH_JOINER

    @staticmethod
    def hide_zero_width(secret_data: bytes, cover_text: str, output_file: str) -> dict:
        """
        Hide binary data as zero-width characters in text.
        Uses 3-character encoding for robustness.
        """
        try:
            # Convert secret to binary string
            binary_str = ''.join(format(byte, '08b') for byte in secret_data)
            
            # Encode binary as zero-width chars (0=ZWSP, 1=ZWNJ)
            encoded = ''
            for bit in binary_str:
                encoded += TXTSteganography.ZERO_WIDTH_ZERO if bit == '0' else TXTSteganography.ZERO_WIDTH_ONE
            
            # Add separator and length header
            length_bits = format(len(secret_data), '032b')
            length_encoded = ''.join(
                TXTSteganography.ZERO_WIDTH_ZERO if b == '0' else TXTSteganography.ZERO_WIDTH_ONE
                for b in length_bits
            )
            
            # Combine: length_header + separator + data + separator
            stego_text = length_encoded + TXTSteganography.ZERO_WIDTH_SEP + encoded + TXTSteganography.ZERO_WIDTH_SEP
            
            # Insert into cover text at random positions (between words)
            lines = cover_text.split('\n')
            if lines:
                # Add to end of first line to preserve readability
                lines[0] += stego_text
            else:
                lines = [stego_text]
            
            result_text = '\n'.join(lines)
            
            # Write output
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(result_text)
            
            return {
                'status': 'success',
                'method': 'zero-width',
                'secret_size': len(secret_data),
                'cover_size': len(cover_text),
                'output_file': output_file,
                'capacity_used': len(encoded) / (len(cover_text) * 8) * 100
            }
        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    @staticmethod
    def extract_zero_width(stego_file: str) -> Tuple[bytes, dict]:
        """Extract binary data from zero-width characters in text"""
        try:
            with open(stego_file, 'r', encoding='utf-8') as f:
                stego_text = f.read()
            
            # Extract length header (32 bits)
            length_bits = ''
            pos = 0
            for _ in range(32):
                if pos < len(stego_text):
                    char = stego_text[pos]
                    if char == TXTSteganography.ZERO_WIDTH_ZERO:
                        length_bits += '0'
                    elif char == TXTSteganography.ZERO_WIDTH_ONE:
                        length_bits += '1'
                    pos += 1
            
            # Convert length
            data_length = int(length_bits, 2) if length_bits else 0
            
            # Skip separator
            if pos < len(stego_text) and stego_text[pos] == TXTSteganography.ZERO_WIDTH_SEP:
                pos += 1
            
            # Extract data bits
            data_bits = ''
            for _ in range(data_length * 8):
                if pos < len(stego_text):
                    char = stego_text[pos]
                    if char == TXTSteganography.ZERO_WIDTH_ZERO:
                        data_bits += '0'
                    elif char == TXTSteganography.ZERO_WIDTH_ONE:
                        data_bits += '1'
                    pos += 1
            
            # Convert bits to bytes
            secret_data = bytes(int(data_bits[i:i+8], 2) for i in range(0, len(data_bits), 8))
            
            return secret_data, {'status': 'success', 'method': 'zero-width', 'size': len(secret_data)}
        except Exception as e:
            return b'', {'status': 'error', 'message': str(e)}

    @staticmethod
    def hide_whitespace(secret_data: bytes, cover_text: str, output_file: str) -> dict:
        """Hide data as trailing whitespace (tab/space) on each line"""
        try:
            # Convert to binary
            binary_str = ''.join(format(byte, '08b') for byte in secret_data)
            
            lines = cover_text.split('\n')
            bit_index = 0
            
            for i in range(len(lines)):
                if bit_index >= len(binary_str):
                    break
                
                # Add to this line
                for _ in range(min(8, len(binary_str) - bit_index)):
                    bit = binary_str[bit_index]
                    lines[i] += '\t' if bit == '1' else ' '
                    bit_index += 1
            
            result_text = '\n'.join(lines)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(result_text)
            
            return {
                'status': 'success',
                'method': 'whitespace',
                'secret_size': len(secret_data),
                'bits_hidden': bit_index
            }
        except Exception as e:
            return {'status': 'error', 'message': str(e)}


class PDFSteganography:
    """PDF steganography using metadata and streams"""

    @staticmethod
    def hide_in_metadata(secret_data: bytes, pdf_path: str, output_path: str) -> dict:
        """Inject secret data into PDF metadata"""
        try:
            from PyPDF2 import PdfReader, PdfWriter
            
            # Read PDF
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            # Copy pages
            for page in reader.pages:
                writer.add_page(page)
            
            # Add metadata with encoded secret
            # Encode secret as hex string
            hex_secret = secret_data.hex()
            
            writer.add_metadata({
                '/Producer': f'Stego-{hex_secret[:50]}',
                '/Creator': f'SteganoSystem-{hex_secret[50:100] if len(hex_secret) > 50 else ""}',
                '/Subject': f'Hidden:{hex_secret[100:] if len(hex_secret) > 100 else ""}',
            })
            
            # Write output
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            return {
                'status': 'success',
                'method': 'metadata',
                'secret_size': len(secret_data),
                'pdf_file': output_path
            }
        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    @staticmethod
    def extract_from_metadata(pdf_path: str) -> Tuple[bytes, dict]:
        """Extract secret data from PDF metadata"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(pdf_path)
            metadata = reader.metadata
            
            if not metadata:
                return b'', {'status': 'error', 'message': 'No metadata found'}
            
            # Try to extract from metadata fields
            hex_secret = ''
            
            # Concatenate all metadata fields
            for key in ['/Producer', '/Creator', '/Subject']:
                if key in metadata:
                    value = metadata[key]
                    if isinstance(value, bytes):
                        value = value.decode('utf-8', errors='ignore')
                    if isinstance(value, str):
                        # Extract hex after prefix
                        if '-' in value:
                            hex_secret += value.split('-', 1)[1]
            
            # Convert hex to bytes
            try:
                secret_data = bytes.fromhex(hex_secret)
                return secret_data, {'status': 'success', 'method': 'metadata', 'size': len(secret_data)}
            except:
                return b'', {'status': 'error', 'message': 'Failed to decode secret from metadata'}
        except Exception as e:
            return b'', {'status': 'error', 'message': str(e)}

    @staticmethod
    def hide_eof(secret_data: bytes, pdf_path: str, output_path: str) -> dict:
        """Append secret data after PDF EOF marker"""
        try:
            # Read original PDF
            with open(pdf_path, 'rb') as f:
                pdf_content = f.read()
            
            # Add marker and secret
            marker = b'STEGO_START:'
            secret_with_marker = marker + secret_data
            
            # Write to output
            with open(output_path, 'wb') as f:
                f.write(pdf_content)
                f.write(b'\n' + secret_with_marker)
            
            return {
                'status': 'success',
                'method': 'eof',
                'secret_size': len(secret_data),
                'original_size': len(pdf_content),
                'output_file': output_path
            }
        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    @staticmethod
    def extract_eof(pdf_path: str) -> Tuple[bytes, dict]:
        """Extract secret data appended after PDF EOF"""
        try:
            with open(pdf_path, 'rb') as f:
                content = f.read()
            
            # Find marker
            marker = b'STEGO_START:'
            if marker in content:
                pos = content.rfind(marker)
                secret_data = content[pos + len(marker):]
                return secret_data, {'status': 'success', 'method': 'eof', 'size': len(secret_data)}
            
            return b'', {'status': 'error', 'message': 'No steganographic marker found'}
        except Exception as e:
            return b'', {'status': 'error', 'message': str(e)}


class DOCXSteganography:
    """DOCX steganography using XML and metadata"""

    @staticmethod
    def hide_in_properties(secret_data: bytes, docx_path: str, output_path: str) -> dict:
        """Hide secret in DOCX core properties with chunking for large data"""
        try:
            from docx import Document
            
            # Open DOCX
            doc = Document(docx_path)
            
            # Access core properties
            core_props = doc.core_properties
            
            # Encode secret as hex
            hex_secret = secret_data.hex()
            
            # Store in multiple properties, each chunk <= 200 chars to be safe
            chunk_size = 200
            chunks = [hex_secret[i:i+chunk_size] for i in range(0, len(hex_secret), chunk_size)]
            
            # Store chunks in different properties
            if len(chunks) > 0:
                core_props.subject = f'Stego:0:{chunks[0]}'  # Chunk 0
            if len(chunks) > 1:
                core_props.keywords = f'Stego:1:{chunks[1]}'  # Chunk 1
            if len(chunks) > 2:
                core_props.comments = f'Stego:2:{chunks[2]}'  # Chunk 2
            if len(chunks) > 3:
                core_props.category = f'Stego:3:{chunks[3]}'  # Chunk 3
            
            # Save
            doc.save(output_path)
            
            return {
                'status': 'success',
                'method': 'properties',
                'secret_size': len(secret_data),
                'chunks': len(chunks),
                'docx_file': output_path
            }
        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    @staticmethod
    def extract_from_properties(docx_path: str) -> Tuple[bytes, dict]:
        """Extract secret from DOCX properties (handles chunked data)"""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            core_props = doc.core_properties
            
            hex_secret = ''
            chunks_found = 0
            
            # Extract from multiple properties (they're chunked)
            property_list = [
                ('subject', core_props.subject),
                ('keywords', core_props.keywords),
                ('comments', core_props.comments),
                ('category', core_props.category),
            ]
            
            for prop_name, prop_value in property_list:
                if prop_value and 'Stego:' in prop_value:
                    # Format is "Stego:N:hexdata"
                    parts = prop_value.split(':', 2)
                    if len(parts) >= 3:
                        hex_secret += parts[2]  # Add chunk
                        chunks_found += 1
            
            # Convert to bytes
            try:
                secret_data = bytes.fromhex(hex_secret)
                return secret_data, {
                    'status': 'success',
                    'method': 'properties',
                    'size': len(secret_data),
                    'chunks': chunks_found
                }
            except:
                return b'', {'status': 'error', 'message': 'Failed to decode secret'}
        except Exception as e:
            return b'', {'status': 'error', 'message': str(e)}

    @staticmethod
    def hide_in_custom_xml(secret_data: bytes, docx_path: str, output_path: str) -> dict:
        """Hide secret in custom XML parts"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            from io import BytesIO
            
            # Read DOCX as ZIP
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Extract all files
                docx_files = {}
                for item in zip_ref.infolist():
                    docx_files[item.filename] = zip_ref.read(item.filename)
            
            # Create custom XML
            hex_secret = secret_data.hex()
            custom_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<customData xmlns="http://stego.system/custom">
  <secret>{hex_secret}</secret>
</customData>'''.encode('utf-8')
            
            # Add to docx_files
            docx_files['customXml/item1.xml'] = custom_xml
            docx_files['customXml/itemProps1.xml'] = b'''<?xml version="1.0" encoding="UTF-8"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" 
            xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
</Properties>'''
            
            # Rewrite DOCX
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for filename, content in docx_files.items():
                    zip_out.writestr(filename, content)
            
            return {
                'status': 'success',
                'method': 'custom-xml',
                'secret_size': len(secret_data),
                'docx_file': output_path
            }
        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    @staticmethod
    def extract_from_custom_xml(docx_path: str) -> Tuple[bytes, dict]:
        """Extract secret from custom XML"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Try to read custom XML
                try:
                    custom_xml_data = zip_ref.read('customXml/item1.xml')
                    root = ET.fromstring(custom_xml_data)
                    
                    # Find secret element
                    for elem in root.iter():
                        if 'secret' in elem.tag.lower():
                            hex_secret = elem.text
                            secret_data = bytes.fromhex(hex_secret)
                            return secret_data, {'status': 'success', 'method': 'custom-xml', 'size': len(secret_data)}
                except:
                    pass
            
            return b'', {'status': 'error', 'message': 'No custom XML secret found'}
        except Exception as e:
            return b'', {'status': 'error', 'message': str(e)}


class DocumentSteganography:
    """Unified document steganography interface"""

    def __init__(self):
        self.txt = TXTSteganography()
        self.pdf = PDFSteganography()
        self.docx = DOCXSteganography()

    def detect_format(self, file_path: str) -> str:
        """Detect document format from extension"""
        if not os.path.exists(file_path):
            return 'unknown'
        
        ext = os.path.splitext(file_path)[1].lower()
        
        # Map extensions to formats
        format_map = {
            '.txt': 'txt',
            '.text': 'txt',
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image',
            '.bmp': 'image',
            '.mp4': 'video',
            '.avi': 'video',
            '.mov': 'video',
            '.wav': 'audio',
            '.mp3': 'audio',
            '.flac': 'audio',
        }
        
        return format_map.get(ext, 'unknown')

    def hide(self, secret_data: bytes, cover_file: str, output_file: str, method: str = 'default') -> dict:
        """Generic hide operation - handles both document and multimedia formats"""
        fmt = self.detect_format(cover_file)
        
        # Handle document cover formats
        if fmt == 'txt':
            with open(cover_file, 'r', encoding='utf-8') as f:
                cover_text = f.read()
            return TXTSteganography.hide_zero_width(secret_data, cover_text, output_file)
        elif fmt == 'pdf':
            return PDFSteganography.hide_in_metadata(secret_data, cover_file, output_file)
        elif fmt == 'docx':
            return DOCXSteganography.hide_in_properties(secret_data, cover_file, output_file)
        
        # Handle multimedia cover formats (image, video, audio)
        # For now, just return error - these require specialized handlers
        elif fmt in ['image', 'video', 'audio']:
            return {
                'status': 'error',
                'message': f'Document-to-{fmt.capitalize()} requires multimedia handler. Use NineConceptsSteganography instead.'
            }
        else:
            return {'status': 'error', 'message': f'Unsupported format: {fmt}'}

    def extract(self, stego_file: str) -> Tuple[bytes, dict]:
        """Generic extract operation"""
        fmt = self.detect_format(stego_file)
        
        if fmt == 'txt':
            return TXTSteganography.extract_zero_width(stego_file)
        elif fmt == 'pdf':
            # Try metadata first, then EOF
            data, info = PDFSteganography.extract_from_metadata(stego_file)
            if data:
                return data, info
            return PDFSteganography.extract_eof(stego_file)
        elif fmt == 'docx':
            return DOCXSteganography.extract_from_properties(stego_file)
        else:
            return b'', {'status': 'error', 'message': f'Unsupported format: {fmt}'}

    def validate_capacity(self, secret_size: int, cover_file: str, method: str = 'default') -> dict:
        """Check if secret can fit in cover"""
        fmt = self.detect_format(cover_file)
        
        try:
            cover_size = os.path.getsize(cover_file)
        except:
            return {'valid': False, 'message': 'Cannot read cover file'}
        
        # Rough capacity estimates (conservative)
        capacity_ratios = {
            'txt': 0.10,      # 10% of file size
            'pdf': 0.15,      # 15% of file size
            'docx': 0.20,     # 20% of file size
        }
        
        max_capacity = cover_size * capacity_ratios.get(fmt, 0.10)
        
        return {
            'valid': secret_size <= max_capacity,
            'secret_size': secret_size,
            'max_capacity': max_capacity,
            'cover_size': cover_size,
            'utilization': (secret_size / max_capacity * 100) if max_capacity > 0 else 0,
            'format': fmt
        }


__all__ = ['DocumentSteganography', 'TXTSteganography', 'PDFSteganography', 'DOCXSteganography']
